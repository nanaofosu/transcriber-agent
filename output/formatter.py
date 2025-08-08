"""Functions for formatting transcription data into various file types.

Each formatter takes either a raw transcript or a list of segments and
returns data ready to be written to a file.  Plain text and markdown
formatters return simple strings, while the DOCX formatter returns a
``python-docx`` Document instance.
"""

from __future__ import annotations

from typing import List, Dict, Any

try:
    from docx import Document  # type: ignore
except ModuleNotFoundError:
    Document = None  # type: ignore


def format_plain_text(text: str) -> str:
    """Return the transcript as plain text.

    Strips any leading and trailing whitespace and ensures that the result
    ends with a single newline.
    """
    return text.strip() + "\n"


def format_markdown(text: str) -> str:
    """Return the transcript as markdown.

    At present this simply returns the text itself.  In future this
    function could add markdown headings or speaker annotations.
    """
    return text.strip() + "\n"


def _seconds_to_timestamp(seconds: float) -> str:
    """Convert a floating‑point second count into an SRT timestamp.

    SRT timestamps take the form ``HH:MM:SS,mmm`` where ``mmm`` are
    milliseconds.  Hours and minutes are zero‑padded to two digits.
    """
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int((seconds - int(seconds)) * 1_000)
    return f"{hours:02}:{minutes:02}:{secs:02},{millis:03}"


def format_srt(segments: List[Dict[str, Any]]) -> str:
    """Convert Whisper segments into an SRT subtitle string.

    Parameters
    ----------
    segments:
        A list of dictionaries as returned by Whisper.  Each dictionary must
        contain ``start``, ``end`` and ``text`` keys.

    Returns
    -------
    str
        The concatenated SRT transcript.
    """
    srt_lines: List[str] = []
    for index, segment in enumerate(segments, start=1):
        start_ts = _seconds_to_timestamp(float(segment.get("start", 0)))
        end_ts = _seconds_to_timestamp(float(segment.get("end", 0)))
        text = segment.get("text", "").strip()
        srt_lines.append(f"{index}\n{start_ts} --> {end_ts}\n{text}\n")
    return "\n".join(srt_lines)


def format_docx(text: str) -> Document:
    """Create a Word document containing the transcript.

    Parameters
    ----------
    text:
        The transcript to embed in the document.

    Returns
    -------
    docx.Document
        A document with the transcript added as a single paragraph.
    """
    if Document is None:
        raise RuntimeError(
            "python-docx is not installed. Please install it to generate DOCX files."
        )
    doc = Document()
    # Add a title to the document for clarity.
    doc.add_heading("Transcription", level=1)
    doc.add_paragraph(text.strip())
    return doc
